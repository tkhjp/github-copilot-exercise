#!/usr/bin/env python
"""Generate DOCX inputs for the Copilot text-vs-image extension.

The DOCX arm is intentionally minimal: each document is a title + the judgment
prompt + the SAME PNG embedded as a picture. That keeps the factual content
identical to the PNG and PPTX arms while testing whether Copilot's docx intake
(MIME handling, implicit OCR/parse, etc.) alters the outcome.

Run:
    python tests/text_vs_image/generate_test_docx.py
    # writes tests/text_vs_image/inputs/{02_ui_change,03_complex_arch}.docx
"""
from __future__ import annotations

from pathlib import Path

import yaml
from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent.parent
IMAGES = ROOT / "images"
INPUTS = ROOT / "inputs"
CASES_YAML = ROOT / "test_cases.yaml"


def _load_case(case_id: str) -> dict:
    data = yaml.safe_load(CASES_YAML.read_text(encoding="utf-8"))
    for c in data["test_cases"]:
        if c["id"] == case_id:
            return c
    raise SystemExit(f"case {case_id} not found in {CASES_YAML}")


def build_docx(case_id: str, png_path: Path, out_path: Path) -> None:
    case = _load_case(case_id)
    title = case.get("title", case_id)
    question = (case.get("question") or "").strip()

    doc = Document()

    # Body default font
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading(title, level=1)
    doc.add_paragraph(question)
    doc.add_paragraph()  # spacing
    doc.add_picture(str(png_path), width=Inches(6.5))

    doc.save(str(out_path))


def main() -> None:
    INPUTS.mkdir(parents=True, exist_ok=True)
    specs = [
        ("tc02_judge", IMAGES / "02_ui_change.png", INPUTS / "02_ui_change.docx"),
        ("tc03_judge", IMAGES / "03_complex_arch.png", INPUTS / "03_complex_arch.docx"),
    ]
    for case_id, png, out in specs:
        if not png.exists():
            raise SystemExit(f"missing PNG: {png}")
        build_docx(case_id, png, out)
        print(f"wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
